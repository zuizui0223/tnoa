#!/usr/bin/env python3
"""Build an anonymous MEE initial-submission DOCX from the canonical source.

Pandoc/citeproc renders citations and the bibliography. ``python-docx`` then adds
single-column submission formatting: double-spaced paragraph styles, continuous
line numbering, page numbering and anonymous core properties. The scientific
manuscript source is not modified.
"""
from __future__ import annotations

import argparse
import subprocess
import sys
import tempfile
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
ASSEMBLER = ROOT / "scripts" / "build_mee_initial_submission_source.py"
SOURCE = ROOT / "submission" / "generated" / "MEE_INITIAL_SUBMISSION_SOURCE.md"
BIB = ROOT / "references.bib"
DEFAULT_OUT = ROOT / "submission" / "generated" / "TNOA_MEE_ANONYMOUS_INITIAL_SUBMISSION.docx"


def fail(message: str) -> None:
    raise SystemExit(f"MEE DOCX build failed: {message}")


def _line_number_section(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    line = OxmlElement("w:lnNumType")
    line.set(qn("w:countBy"), "1")
    line.set(qn("w:start"), "1")
    line.set(qn("w:restart"), "continuous")
    sect_pr.append(line)

    existing_pg = sect_pr.find(qn("w:pgNumType"))
    if existing_pg is not None:
        sect_pr.remove(existing_pg)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "1")
    sect_pr.append(pg)


def _page_number_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def _format_docx(input_path: Path, output_path: Path) -> None:
    doc = Document(str(input_path))

    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style.paragraph_format.line_spacing = 2.0
            style.paragraph_format.space_after = Pt(0)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        _line_number_section(section)
        _page_number_footer(section)

    props = doc.core_properties
    props.author = "Anonymous"
    props.last_modified_by = "Anonymous"
    props.title = "Preserving unresolved observations in automated ecological sensing: a target–nuisance–observability framework"
    props.subject = "Anonymous initial submission to Methods in Ecology and Evolution"
    props.comments = "Generated reproducibly from the audited anonymous MEE source."

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csl", type=Path, required=True, help="Pinned independent CSL style used by Pandoc citeproc")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()

    for path in (ASSEMBLER, BIB, args.csl):
        if not path.is_file():
            fail(f"missing required source: {path}")

    subprocess.run([sys.executable, str(ASSEMBLER)], check=True)
    if not SOURCE.is_file():
        fail("canonical MEE source was not generated")

    with tempfile.TemporaryDirectory(prefix="tnoa-docx-") as tmp:
        raw = Path(tmp) / "pandoc.docx"
        extra = [
            "--citeproc",
            f"--bibliography={BIB}",
            f"--csl={args.csl}",
            "--metadata=reference-section-title:References",
            "--standalone",
        ]
        result = pypandoc.convert_file(
            str(SOURCE),
            to="docx",
            format="markdown",
            outputfile=str(raw),
            extra_args=extra,
        )
        if result not in ("", None):
            fail(f"unexpected Pandoc return value: {result!r}")
        if not raw.is_file() or raw.stat().st_size < 10000:
            fail("Pandoc did not produce a plausible DOCX")
        _format_docx(raw, args.output)

    print(f"Built anonymous MEE DOCX: {args.output}")


if __name__ == "__main__":
    main()
